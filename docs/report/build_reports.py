#!/usr/bin/env python3
"""Build expanded NEU internship summary reports (~20 pages ZH + EN)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent
FIG = ROOT / "figures"
FIG_EN = ROOT / "figures" / "en"
OUT_DIR = ROOT


def set_run_font(run, name_cn="宋体", name_en="Times New Roman", size=12, bold=False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name_cn)


def add_para(doc, text, *, size=12, bold=False, align="left", first_line=True, space_after=6):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line and align == "left":
        pf.first_line_indent = Cm(0.74)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading_cn(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, name_cn="黑体", name_en="Arial", size=sizes.get(level, 12), bold=True)
    return p


def add_heading_en(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, name_cn="Arial", name_en="Arial", size=sizes.get(level, 12), bold=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, name_cn="楷体", name_en="Times New Roman", size=10.5)


def add_figure(doc, path: Path, caption: str, width_cm=14.2):
    if not path.exists():
        add_para(doc, f"[Missing figure: {path.name}]", first_line=False, align="center")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(text)
    set_run_font(run, name_cn="Consolas", name_en="Consolas", size=9)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, name_cn="黑体", name_en="Arial", size=10, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
    doc.add_paragraph()


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top/left/bottom/right -> {'sz': '12', 'val': 'single', 'color': '000000'}."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge not in kwargs:
            continue
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tcBorders.append(element)
        for k, v in kwargs[edge].items():
            element.set(qn(f"w:{k}"), str(v))


def set_run_cell_text(cell, text, *, size=12, bold=False, align="center", name_cn="宋体", name_en="Times New Roman"):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, name_cn=name_cn, name_en=name_en, size=size, bold=bold)
    return p


def set_paragraph_bottom_border(paragraph, sz="12", color="000000"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def setup_page_header(section, left: str, right: str, *, en: bool = False):
    """Left/right header with a full-width bottom rule; cover uses different first page (no header)."""
    section.different_first_page_header_footer = True

    def _clear_header(hdr):
        # keep one empty paragraph
        for p in hdr.paragraphs:
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            p.text = ""
        while len(hdr.paragraphs) > 1:
            el = hdr.paragraphs[-1]._element
            el.getparent().remove(el)

    # empty first-page header (cover)
    _clear_header(section.first_page_header)

    header = section.header
    header.is_linked_to_previous = False
    _clear_header(header)
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    usable = section.page_width - section.left_margin - section.right_margin
    pf.tab_stops.add_tab_stop(usable, WD_TAB_ALIGNMENT.RIGHT)
    run_l = p.add_run(left)
    if en:
        set_run_font(run_l, name_cn="Arial", name_en="Times New Roman", size=10.5)
    else:
        set_run_font(run_l, name_cn="宋体", name_en="Times New Roman", size=10.5)
    p.add_run("\t")
    run_r = p.add_run(right)
    if en:
        set_run_font(run_r, name_cn="Arial", name_en="Times New Roman", size=10.5)
    else:
        set_run_font(run_r, name_cn="宋体", name_en="Times New Roman", size=10.5)
    set_paragraph_bottom_border(p, sz="12")
    section.header_distance = Cm(1.5)


def add_cover_field(doc, label: str, value: str, *, en: bool = False):
    """Template-like fill-in line: label + underlined value (left-aligned form block)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(14)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Cm(4.2)
    run_l = p.add_run(label)
    set_run_font(
        run_l,
        name_cn="宋体" if not en else "Arial",
        name_en="Times New Roman" if not en else "Arial",
        size=14,
    )
    display = value if value else "____________________"
    run_v = p.add_run(display)
    set_run_font(
        run_v,
        name_cn="宋体" if not en else "Arial",
        name_en="Times New Roman" if not en else "Arial",
        size=14,
    )
    run_v.underline = True
    return p


def build_review_box(doc, *, en: bool = False):
    """Cover evaluation box matching the college template layout."""
    if en:
        label_score = "Reviewer\nScore"
        h_item = "Scoring Item"
        h_pts = "Score (100)"
        item1 = "Substantial workload; solutions provided and limitations identified"
        item2 = "Standardized format; content comprehensive and clear"
        label_cmt = "Reviewer\nComments"
        sign = "Signature:"
    else:
        label_score = "评阅人评分"
        h_item = "评分项"
        h_pts = "分数（百分制）"
        item1 = "工作量饱满，给出解决方案并指出局限性"
        item2 = "格式规范，内容全面、清晰"
        label_cmt = "评阅人评语"
        sign = "签字："

    # 4 rows x 3 cols; then merge for comments row
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # widths roughly: label | item | score
    widths = (Cm(2.6), Cm(10.2), Cm(2.8))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # header row of scoring block
    set_run_cell_text(table.cell(0, 0), label_score, size=11, bold=True, name_cn="宋体" if not en else "Arial", name_en="Arial" if en else "Times New Roman")
    set_run_cell_text(table.cell(0, 1), h_item, size=11, bold=True, name_cn="宋体" if not en else "Arial", name_en="Arial" if en else "Times New Roman")
    set_run_cell_text(table.cell(0, 2), h_pts, size=11, bold=True, name_cn="宋体" if not en else "Arial", name_en="Arial" if en else "Times New Roman")

    set_run_cell_text(table.cell(1, 1), item1, size=11, align="left", name_cn="宋体" if not en else "Arial", name_en="Arial" if en else "Times New Roman")
    set_run_cell_text(table.cell(1, 2), "", size=11)
    set_run_cell_text(table.cell(2, 1), item2, size=11, align="left", name_cn="宋体" if not en else "Arial", name_en="Arial" if en else "Times New Roman")
    set_run_cell_text(table.cell(2, 2), "", size=11)

    # merge left column rows 0-2 for 评阅人评分
    a = table.cell(0, 0)
    a.merge(table.cell(2, 0))
    set_run_cell_text(a, label_score, size=11, bold=True, name_cn="宋体" if not en else "Arial", name_en="Arial" if en else "Times New Roman")

    # comments row: merge col1+col2
    c0 = table.cell(3, 0)
    c1 = table.cell(3, 1)
    c2 = table.cell(3, 2)
    c1.merge(c2)
    set_run_cell_text(c0, label_cmt, size=11, bold=True, name_cn="宋体" if not en else "Arial", name_en="Arial" if en else "Times New Roman")
    # tall comments cell with signature at bottom-right
    c1.text = ""
    # blank lines
    for _ in range(5):
        bp = c1.add_paragraph()
        bp.paragraph_format.space_after = Pt(8)
    sp = c1.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sp.add_run(sign)
    set_run_font(run, name_cn="宋体" if not en else "Arial", name_en="Arial" if en else "Times New Roman", size=11)
    # force row height for comments
    tr = table.rows[3]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), "1800")
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)

    doc.add_paragraph()
    return table


def cover_zh(doc):
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("东北大学软件学院")
    set_run_font(run, name_cn="黑体", name_en="Arial", size=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run("学生实训总结报告")
    set_run_font(run, name_cn="黑体", name_en="Arial", size=26, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    for label, value in [
        ("专    业：", "软件工程（软件英才班）"),
        ("班    级：", "软英2301"),
        ("学    号：", "20236742"),
        ("姓    名：", "王瑞琦"),
        ("实训基地：", "东北大学软件学院企业项目实训"),
        ("企业指导教师：", "【请填写】"),
    ]:
        add_cover_field(doc, label, value)

    for _ in range(2):
        doc.add_paragraph()

    build_review_box(doc, en=False)

    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("年　　月　　日")
    set_run_font(run, name_cn="宋体", name_en="Times New Roman", size=12)
    doc.add_page_break()


def cover_en(doc):
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("School of Software, Northeastern University")
    set_run_font(run, name_cn="Arial", name_en="Arial", size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("Student Internship Summary Report")
    set_run_font(run, name_cn="Arial", name_en="Arial", size=22, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    for label, value in [
        ("Major: ", "Software Engineering (Elite Class)"),
        ("Class: ", "SE-EN 2301"),
        ("Student ID: ", "20236742"),
        ("Name: ", "Wang Ruiqi"),
        ("Internship Base: ", "NEU Software College Enterprise Project Training"),
        ("Industry Mentor: ", "[To be filled]"),
    ]:
        add_cover_field(doc, label, value, en=True)

    for _ in range(2):
        doc.add_paragraph()

    build_review_box(doc, en=True)

    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Year　　Month　　Day")
    set_run_font(run, name_cn="Arial", name_en="Arial", size=12)
    doc.add_page_break()


def build_zh():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.6)
    setup_page_header(sec, "东北大学软件学院", "学生实训总结报告", en=False)
    cover_zh(doc)

    add_heading_cn(doc, "说    明", 1)
    for t in [
        "实训结束之前，每个学生都必须认真撰写《实训总结报告》。通过撰写实训总结报告，系统地回顾和总结实训的全过程，将实践性教学的感性认识升华到一定的理论高度，从而提高实训教学效果。",
        "实训总结报告要求条理清晰，内容详尽，数据准确。字数不少于5000字。实训总结报告的质量反映了实训的质量，它是实训成绩评定的主要依据之一。应在规定时间内将此报告上交。不交实训报告者不得参加实训成绩评定。",
        "封面中的“实训基地/实训单位”须写单位全名，并填写企业指导教师姓名（本封面“企业指导教师”处暂留【请填写】，定稿前补全）。",
        "本报告围绕 C 组医学影像标注与数据管理平台中本人实际完成的工作展开，按学院模板组织“前言—实训内容（概述、相关技术、系统分析、系统设计、系统实现、系统测试）—总结”。文中给出架构图、用例图、活动图、类/对象图、状态图、时序图、协作图、程序流程图与关键代码；带“预留/待替换”字样的界面图为占位，本人后续将补入真实运行截图。",
        "本中文版为提交用主稿；另附英文版便于双语归档。字数满足学院不少于五千字的要求。",
    ]:
        add_para(doc, t)
    doc.add_page_break()

    add_heading_cn(doc, "目  录", 1)
    for line in [
        "1  前言",
        "    1.1  实训背景",
        "    1.2  实训环境",
        "    1.3  实训过程",
        "2  实训内容",
        "    2.1  概述",
        "    2.2  相关技术",
        "    2.3  系统分析",
        "        2.3.1  业务分析",
        "        2.3.2  功能分析（用例图、活动图）",
        "        2.3.3  可行性与范围分析",
        "    2.4  系统设计",
        "        2.4.1  总体设计（架构与模块）",
        "        2.4.2  详细设计（类/对象、动态模型、算法）",
        "        2.4.3  数据库设计",
        "    2.5  系统实现",
        "    2.6  系统测试",
        "        2.6.1  测试方案设计",
        "        2.6.2  测试用例与结果",
        "        2.6.3  测试结论与评价",
        "3  总结",
        "    3.1  实训体会",
        "    3.2  其它意见",
        "参考文献与附录说明",
    ]:
        add_para(doc, line, first_line=False, space_after=2)
    doc.add_page_break()

    # ========== 1 ==========
    add_heading_cn(doc, "1  前言", 1)
    add_heading_cn(doc, "1.1  实训背景", 2)
    for t in [
        "这次企业项目实训安排在本科高年级阶段。学院的意图很清楚：不再满足于课程大作业式的“单机可演示页面”，而是让学生在接近真实协作的条件下，把一个医学影像相关系统从文档推进到可运行、可联调、可答辩的半成品。我所在的 C 组题目是医学影像标注与数据管理平台，GitHub 仓库名为 C_Group_Medical_Annotation。",
        "课题背景来自医学影像工作中的现实矛盾。CT 等三维影像的人工勾画耗时长、版本多、回退难；纯自动分割又很难在一次推理后直接达到临床可用。于是平台需要把“AI 粗标—人工精修—审核归档—导出训练数据”做成闭环。对软件工程学生而言，这正好覆盖需求建模、接口设计、数据持久化、人机交互、质量保障与团队协作等多类能力。",
        "组内按平台工程与 AI 算法两条线分工。本人主要负责平台侧：后端 API、数据库、标注前端、三维可视化、权限与审核工作流，以及和 AI 推理服务的对接面。队友侧重模型训练、损失函数、Dataset 管线与推理脚本。双方共用同一仓库、统一文件命名与 /api 契约。对我个人，实训目标可以概括为：走完软件工程完整链条；真正处理 DICOM/NIfTI/Mask 这类“脏数据”；在队友进度不同步时，仍能用契约把系统接住。",
        "需要提前说明的是：本项目属于教学实训向的工程系统，不是已取证的医疗器械软件。报告中写到的“模拟手术”用于检验三维交互与 ROI 数据闭环，不能等同于临床手术规划系统。把边界写清楚，是工程文档应有的诚实态度。",
    ]:
        add_para(doc, t)

    add_heading_cn(doc, "1.2  实训环境", 2)
    for t in [
        "实训基地挂靠东北大学软件学院企业项目实训，以 C 组医学影像标注平台为载体。日常开发环境为 macOS；后端 Python 3 + FastAPI，本机 uvicorn 监听 127.0.0.1:8000；数据层使用 SQLite，表结构由 database/schema.sql 与运行时 ensure_*_schema 共同维护；前端主路径是 frontend/ 下的原生 HTML/CSS/JS 工作台，同时保留 web/ 下 React 相关资源；三维渲染使用 VTK.js，并保留 WebGL2 体绘制路径；手势交互接入浏览器摄像头与 MediaPipe Hands；AI 侧对接 nnU-Net、TotalSegmentator、DeepEdit 等（本机曾配置 TOTALSEG_PYTHON=/opt/miniconda3/bin/python）。",
        "版本管理使用 GitHub，核心分支包括 feature-a（本人）、feature-b（队友）、dev 与 main。大体积影像、模型权重与真实病人数据不进入仓库，只提交代码、文档、配置与 split manifest。演示账号在后端启动时种子写入，例如 annotator / reviewer / admin 三类角色，便于权限联调。",
    ]:
        add_para(doc, t)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["操作系统", "macOS（darwin）"],
            ["后端", "FastAPI + uvicorn，端口 8000"],
            ["数据库", "SQLite（schema.sql + 运行时迁移）"],
            ["前端", "frontend/ 原生工作台；web/ React 资源"],
            ["三维/手势", "VTK.js、WebGL2、MediaPipe Hands"],
            ["AI 对接", "nnU-Net / TotalSeg / DeepEdit / platform_unet"],
            ["协作", "GitHub：feature-a / feature-b / dev / main"],
        ],
    )
    add_para(
        doc,
        "所在“部门”实质上就是平台工程组。我的日常工作包括接口设计与实现、标注台交互、体数据与 Mask 服务、手势与模拟手术模块、与负责 AI 的队友联调预测写回、文档维护与分支合并。企业指导教师姓名请在封面【请填写】处补全。",
    )

    add_heading_cn(doc, "1.3  实训过程", 2)
    for t in [
        "实训过程大致分为四个阶段，阶段之间不是截然断开，而是“主路径跑通后再加深”。第一阶段聚焦标准与骨架：完成数据流与文件命名标准、ER 草案、API 契约和 UI 线框，把 cases / images / masks / versions 跑通，前端能够选病例、看切片、保存 Mask。没有这一步，后面所有三维和 AI 都没有落点。",
        "第二阶段补齐工作流与可视化：加入 JWT 登录、任务分配、提交/审核/驳回；接入 VTK 三维与多平面重建（MPR），并支持 MIP/MinIP 投影；上传从单文件扩展到拖拽与多 DICOM 序列导入。这一阶段的目标是让“标注员的一天”在系统里走得通。",
        "第三阶段加深人机协同：对接真实 AI predict，默认拒绝 silent HU baseline 伪成功；支持多标签标注、版本质量对比与 Dataset 导出；引入 MediaPipe 双手势，完成三维导航与器官点选。队友的多器官 nnU-Net / TotalSeg 能力也在这一阶段通过平台接口被“接进来”。",
        "第四阶段是结题冲刺：把手势面板从遮挡 3D 的浮层，改成插在三维视图与 MPR 之间的文档流 dock；补齐 VTK 多器官分色表面网格（per_label Marching Cubes）；实现模拟手术“选器官→确认长方体 ROI→切割”硬门槛；将 surgery_results（含器官名称与颜色）写入数据库；最后把 feature-a 合并进 main/dev 并推送 GitHub。",
        "过程中有几处返工值得单独记录。其一，联调节奏：队友的数据规范在变，导出字段若写死就会对不上，后来更强调 manifest 与 label 映射可配置。其二，演示数据坑：单层 DCM 无法支撑可靠 3D/手势，于是在入口增加层数与服务健康预检。其三，交互返工：浮层手势面板挡住三维，改成 dock 后观感朴素，但更符合“工具服务于观察”。这些返工看起来不华丽，却很像真实项目里的日常。",
    ]:
        add_para(doc, t)
    add_figure(doc, FIG / "fig3_annotation_flow.png", "图1-1 标注主流程（实训各阶段工作的主干）")
    add_figure(doc, FIG / "fig8_review_flow.png", "图1-2 审核工作流活动图（简化）")

    # ========== 2 ==========
    add_heading_cn(doc, "2  实训内容", 1)
    add_heading_cn(doc, "2.1  概述", 2)
    for t in [
        "项目可表述为“C 组医学影像标注与数据管理平台”。系统面向标注员、审核员和管理员，支持 CT 等影像上传、二维多标签勾画、三维浏览与手势交互、AI 辅助分割写回、版本对比与审核、训练数据导出，以及面向演示的模拟手术 ROI 记录。本人工作集中在平台工程侧：不负责从零训练一个新的 SOTA 网络，但必须让平台“诚实”地调用模型——没有权重就明确失败，而不是悄悄返回一张看起来像分割的伪 Mask。",
        "按模块归纳，我完成或主导的内容包括：FastAPI 后端与 SQLite 仓储；认证、任务与审核 API；影像上传、切片与体数据接口；Mask CRUD、版本、对比与导出；标注工作台前端与工具链；VTK/WebGL2 三维与 MPR；MediaPipe 手势控制与器官聚焦；模拟手术 ROI 与器官信息持久化；与队友的 predict / DeepEdit / TotalSeg 联调面；Git 协作、文档与分支合并发布。",
    ]:
        add_para(doc, t)
    add_figure(doc, FIG / "fig1_architecture.png", "图2-1 系统总体架构")
    add_figure(doc, FIG / "fig7_modules.png", "图2-2 本人负责的功能模块分解")
    add_figure(doc, FIG / "fig_placeholder_ui.png", "图2-3 平台登录界面（预留，待替换为实际运行截图）")
    add_figure(doc, FIG / "fig_placeholder_ui.png", "图2-4 标注工作台总览（预留，待替换）")

    add_heading_cn(doc, "2.2  相关技术", 2)
    add_heading_cn(doc, "2.2.1 理论基础", 3)
    for t in [
        "医学影像标注平台首先涉及医学影像数据表示。DICOM 序列按层组织，NIfTI/NRRD 更常用于科研管线中的体数据与标签。标注结果以 Mask 体数据或多边形/画笔轨迹表达；多标签场景下同一体素网格可编码不同器官 ID。评价分割常用 Dice、IoU 等重叠度指标，平台提供 compare 接口做版本间粗评，但不能替代专家审核。",
        "三维可视化方面，体绘制与表面绘制是两条常见路线。体绘制适合观察密度分布，表面网格（如 Marching Cubes）更适合看清器官边界并做交互选中。MPR（多平面重建）把任意轴的切片与三维联动，是放射科阅读的基本习惯，因此平台把轴位/冠状/矢状与 3D 放在同一工作流里。",
        "人机协同标注的基本假设是：模型提供初值，人负责纠错与确认。因此系统必须同时支持“AI 写回”和“人工覆盖”，并用版本把来源记录下来。手势交互属于增强通道，其理论依据是降低三维场景中的鼠标负担，但必须处理识别噪声与误触发。",
    ]:
        add_para(doc, t)
    add_heading_cn(doc, "2.2.2 开发方法、工具与环境", 3)
    for t in [
        "方法上我们采用迭代增量，而不是一次性抛出完整设计。先冻结 API 与数据命名，再打通纵向切片（上传—浏览—保存），随后按答辩与联调压力增加三维、手势、手术 ROI 等能力。文档（docs/01–16）充当组内说明书，减少口头约定丢失。",
        "工具链方面：后端 FastAPI + Pydantic；前端原生 JS（演示依赖少）与部分 React；三维 VTK.js / WebGL2；手势 MediaPipe；AI 侧 Python 生态（PyTorch/nnU-Net/TotalSeg）；协作 Git + GitHub。本地用 uvicorn 热重载开发，关键脚本用 node --check 与 py_compile 做语法门禁。",
    ]:
        add_para(doc, t)
    add_table(
        doc,
        ["类别", "选型", "选用理由（简述）"],
        [
            ["Web API", "FastAPI", "类型注解、OpenAPI、迭代快"],
            ["DB", "SQLite", "部署轻、足够支撑实训规模"],
            ["前端主路径", "原生 JS 工作台", "演示依赖少、易联调"],
            ["三维", "VTK.js + WebGL2", "表面/体绘制可切换"],
            ["手势", "MediaPipe Hands", "浏览器端可用、接入快"],
            ["AI", "nnU-Net/TotalSeg/DeepEdit", "与队友 AI 能力对齐"],
        ],
    )

    add_heading_cn(doc, "2.3  系统分析", 2)
    add_para(
        doc,
        "系统分析从“整个项目要解决什么问题”落到“本人工作边界内要交付什么”。下面从业务、功能与可行性三方面展开，并用用例图、活动图描述关键交互。",
    )
    add_heading_cn(doc, "2.3.1  业务分析", 3)
    for t in [
        "业务上，一条病例从进入系统到变成可训练样本，通常经历：上传影像 → 生成 case/image 记录 → 标注员打开工作台浏览并勾画（可先 AI 预测）→ 保存 mask 与 version → 提交审核 → 审核通过后导出 Dataset。中间要处理多器官标签、版本并存、失败可回退、权限隔离等问题。",
        "对三维场景，用户还希望旋转观察、MPR 对照、必要时用少接触鼠标的方式操作，以及在演示场景下对选中器官记录“假手术”ROI。这些需求并不替代临床手术规划，但能检验交互链路与数据能否闭环保存。若闭环断在“只能看不能存”或“存了不知道切的是谁”，系统价值会大打折扣。",
        "从涉众看，标注员关心效率与可撤销；审核员关心差异与责任追溯；管理员关心任务分派与账号；AI 同学关心导出格式是否稳定。本人负责的平台必须同时回应这几类诉求，而不是只优化某一个界面动画。",
    ]:
        add_para(doc, t)
    add_heading_cn(doc, "2.3.2  功能分析（用例图、活动图）", 3)
    add_para(
        doc,
        "角色与用例见图2-5。标注员关注上传、浏览、标注、AI、导出准备；审核员关注任务队列与通过/驳回；管理员关注用户与全局配置。实现时把权限边界写进 API：例如创建任务需 admin/reviewer，提交与审核接口校验角色，避免“前端藏按钮”冒充安全。",
    )
    add_figure(doc, FIG / "fig2_usecase.png", "图2-5 核心用例图（平台侧）")
    add_table(
        doc,
        ["用例", "主要角色", "前置条件", "后置结果"],
        [
            ["登录", "全体", "账号存在", "获得 JWT 与角色"],
            ["上传 CT", "标注员/管理员", "已登录", "生成 case/image"],
            ["二维标注保存", "标注员", "打开影像", "mask/version 可查"],
            ["AI 预测", "标注员", "模型可用", "写回 AI mask 或明确失败"],
            ["三维/手势浏览", "标注员", "层数足够", "可旋转/点选器官"],
            ["模拟手术 ROI", "标注员", "已选器官并确认 ROI", "surgery_results 入库"],
            ["审核通过/驳回", "审核员", "已提交", "状态变更并可追溯"],
            ["导出 Dataset", "标注员/管理员", "存在可用版本", "生成训练目录与 manifest"],
        ],
    )
    add_para(
        doc,
        "活动图用于描述跨角色、跨步骤的业务过程。图2-5a 给出标注主活动（从上传到导出）；图2-5b 给出审核活动（提交—通过/驳回—再编辑）。这两条活动与后文状态机、时序图相互对应。",
    )
    add_figure(doc, FIG / "fig3_annotation_flow.png", "图2-5a 标注主流程活动图")
    add_figure(doc, FIG / "fig8_review_flow.png", "图2-5b 审核工作流活动图")
    add_heading_cn(doc, "2.3.3  可行性与范围分析", 3)
    for t in [
        "技术可行：组员具备 Web 与 Python 基础，开源组件齐全。经济与进度可行：SQLite 与本地文件降低部署成本，演示用公开多层 CT（如 Case0002–0004），明确拒绝单层无法做 3D 的序列。",
        "操作可行：浏览器端即可完成登录、标注、三维与手势；AI 推理在本机或约定 Python 环境下执行。法律与伦理上，本系统仅使用教学/公开样例数据，不采集真实患者隐私信息。",
        "范围上我们承认若干局限：平台自研 U-Net 偏教学向 2.5D；高精度分割仍依赖 TotalSeg/nnU-Net；模拟手术不是临床器械规划；手势受光照与摄像头影响大。把局限写进分析，是为了避免答辩时把“能演示”说成“已临床可用”。",
    ]:
        add_para(doc, t)
    add_table(
        doc,
        ["优先级", "能力", "说明"],
        [
            ["P0", "上传/浏览/标注/保存/审核/导出", "没有它们不能称为标注系统"],
            ["P0", "AI 预测对接与诚实失败", "避免伪成功污染数据"],
            ["P1", "VTK 三维、MPR、MIP/MinIP", "提升理解与质检效率"],
            ["P1", "手势交互与器官聚焦", "降低三维操作负担"],
            ["P2", "模拟手术 ROI 入库", "演示闭环与扩展点"],
        ],
    )

    add_heading_cn(doc, "2.4  系统设计", 2)
    add_heading_cn(doc, "2.4.1  总体设计（架构与模块）", 3)
    for t in [
        "总体结构按“表现层—API—服务—数据/文件”分层（见图2-1）。表现层负责交互与渲染；API 层做鉴权、参数校验与路由；服务层承载业务规则；数据层包含 SQLite 与影像/Mask 文件。前端不直连数据库，AI 也不绕过平台私自改正式版本，这是为了保证审计与回滚。",
        "功能模块见图2-2：认证与审核、影像服务、Mask 与版本、AI 代理、三维渲染、手势交互、模拟手术、前端工作台。仓库目录上，backend/、frontend/、database/、docs/ 由本人主责，ai/ 由队友主责，dataset/ 双方共享约定。",
    ]:
        add_para(doc, t)
    add_heading_cn(doc, "2.4.2  详细设计（类/对象、动态模型、算法）", 3)
    add_para(
        doc,
        "围绕本人工作内容，详细设计从类/对象、动态模型与算法三方面展开。后端不以重量级 ORM 类层次为主，而以“路由 Schema + Service 函数 + 领域记录”组织；前端以模块对象协作。",
    )
    add_figure(doc, FIG / "fig11_class.png", "图2-6a 核心类与对象设计（节选）")
    add_para(
        doc,
        "动态模型包括：状态图（审核状态迁移 + 模拟手术三步门槛，见图2-6b）、时序图（保存手术 ROI，见图2-8）、协作图（AI 预测写回，见图2-6c）。状态转换条件分别是：选中有效 label；用户确认长方体 ROI；进入可切割并累积刀痕；审核侧则是标注中→已提交→已通过/已驳回。",
    )
    add_figure(doc, FIG / "fig12_state.png", "图2-6b 状态图（审核 + 手术模式）")
    add_figure(doc, FIG / "fig6_sequence.png", "图2-8 保存手术 ROI 时序图（简化）")
    add_figure(doc, FIG / "fig14_collab.png", "图2-6c AI 预测写回协作图（简化）")
    add_figure(doc, FIG / "fig4_surgery_flow.png", "图2-7 模拟手术 ROI 三步流程")
    add_figure(doc, FIG / "fig9_gesture_map.png", "图2-9 手势到三维操作的映射（示意）")
    for t in [
        "算法设计要点：（1）器官字段解析算法——请求体 organ 对象优先，缺省回退 label 目录补全 name/display_name/color，保证库中可读；（2）ROI 合法性——cuboid_max≥cuboid_min，label_id>0，case/image 归属一致；（3）表面网格 per_label 分层——按标签拆分等值面，避免多器官同色；（4）版本 compare——Dice/IoU 粗评，不替代审核。",
        "接口统一前缀 /api。典型路径包括 /api/auth/login、/api/cases、/api/upload、/api/image/{id}/slice/...、/api/save_mask、/api/ai/predict、/api/export、/api/surgery_results。三维网格支持 /api/mask/{id}/surface-mesh?per_label=true。",
    ]:
        add_para(doc, t)
    add_heading_cn(doc, "2.4.3  数据库设计", 3)
    for t in [
        "核心实体包括 users、cases、images、annotations、masks、versions、models、datasets/tasks，以及后期增加的 surgery_results。cases 是病例中心；images 归属病例；masks 关联图像与标签；versions 把一次可追溯的标注状态固化下来；surgery_results 记录模拟手术 ROI。",
        "surgery_results 除 cuboid_min/max、cut_planes、carved_voxels 外，增加 organ_name、organ_display_name、organ_color、organ_json，保证记录能回答“切的是哪个器官”。对旧库，ensure_surgery_schema 通过 PRAGMA table_info + ALTER TABLE 增量加列，避免演示库推倒重来。",
    ]:
        add_para(doc, t)
    add_figure(doc, FIG / "fig5_er.png", "图2-10 核心数据实体关系（节选）")
    add_table(
        doc,
        ["表", "关键字段", "作用"],
        [
            ["cases", "case_id, status, patient_id", "病例主数据与状态"],
            ["images", "image_id, case_id, path, shape", "影像路径与尺寸元信息"],
            ["masks", "mask_id, image_id, label_id, path", "分割结果存储"],
            ["versions", "version_id, case_id, mask_id, tag", "版本与来源追溯"],
            ["users/tasks", "role, assignee, deadline", "权限与任务"],
            ["surgery_results", "organ_*, cuboid_*, cut_planes", "手术 ROI 与器官信息"],
        ],
    )

    add_heading_cn(doc, "2.5  系统实现", 2)
    add_para(
        doc,
        "实现阶段按“先打通主路径，再补体验与边界”推进。下面从编码角度论述本人主要模块，并通过程序流程图、关键代码与界面（预留）直观说明。",
    )
    add_heading_cn(doc, "2.5.1  后端、认证与数据访问", 3)
    for t in [
        "backend/app/main.py 注册路由；业务落在 services/*。SQLite 连接统一管理。认证采用 JWT：登录成功后前端在 Header 携带 Bearer Token，受保护写接口校验角色。任务与审核接口把 case 状态迁移与审计日志串起来，便于追溯“谁在什么时候通过了什么”。",
        "种子账号方便本地联调，但正式使用仍应改密并限制外网暴露。我在实现里把角色校验放在服务端，是因为前端隐藏按钮挡不住直接调 API。实训里这一点经常被忽略，却是权限设计的底线。",
        "以手术结果为例，保存时先校验 case/image 归属与 label_id，再规范化 cut_planes，最后解析器官字段。请求体可带 organ 对象；缺省则回退 label 目录补全名称与颜色，避免库里只有整数 label。",
    ]:
        add_para(doc, t)
    add_figure(doc, FIG / "fig13_program_flow.png", "图2-11 保存手术 ROI 程序流程图")
    add_code(
        doc,
        "POST /api/surgery_results\n"
        "body: case_id, image_id, mask_id?, label_id,\n"
        "      organ_name?, organ_display_name?, organ_color?, organ?,\n"
        "      roi_margin_pct, knife_radius, cuboid_min, cuboid_max,\n"
        "      cut_planes[], carved_voxels, note?",
    )
    add_code(
        doc,
        "# surgery_service.py（关键片段，语义化）\n"
        "def save_surgery_result(request, user=None):\n"
        "    ensure_surgery_schema()\n"
        "    case = get_record('cases', 'case_id', request.case_id)\n"
        "    if case is None: raise HTTPException(404, ...)\n"
        "    image = get_record('images', 'image_id', request.image_id)\n"
        "    if image['case_id'] != request.case_id: raise HTTPException(422, ...)\n"
        "    if request.label_id <= 0: raise HTTPException(422, ...)\n"
        "    organ = _resolve_organ_fields(request)  # 请求优先，目录回退\n"
        "    # INSERT surgery_results(... organ_name/display/color ...)\n"
        "    return SaveSurgeryResultResponse(...)",
    )
    add_table(
        doc,
        ["字段", "含义", "来源"],
        [
            ["label_id", "器官标签编号", "当前选中器官"],
            ["organ_name", "器官内部名", "请求或标签目录"],
            ["organ_display_name", "界面显示名", "请求或目录"],
            ["organ_color", "可视化颜色", "请求或目录"],
            ["cuboid_min/max", "长方体对角点", "确认后的 ROI"],
            ["cut_planes", "刀痕平面列表", "切割过程累积"],
            ["carved_voxels", "雕刻体素数（粗计）", "前端统计回传"],
        ],
    )
    add_heading_cn(doc, "2.5.2  影像、Mask 与导出", 3)
    for t in [
        "上传支持 NIfTI/常见影像包与多 DICOM 序列，后端负责落盘、建 image 记录并提供切片 PNG 与 volume 元信息。Mask 保存按 label_id 写入，支持更新与删除；版本机制把 manual / ai / final 等语义分开。compare 接口计算 Dice/IoU，用于快速查看 AI 与人工差异。",
        "导出 Dataset 时，可选择 materialize 把多器官 mask 合并为多类 labels，并写出 dataset.json 与划分 manifest，供队友训练读取。这里的关键是“导出可重复”：同样的 case 集合与版本标签，应得到结构稳定的目录树。",
    ]:
        add_para(doc, t)
    add_figure(doc, FIG / "fig_placeholder_ui.png", "图2-12 版本列表与对比面板（预留，待替换）")
    add_heading_cn(doc, "2.5.3  标注工作台前端", 3)
    for t in [
        "frontend/app.js 是工作台中枢：病例列表、工具栏、AI 按钮、版本面板、保存与导出都从这里调度。多标签下画笔/橡皮按当前 label_id 写入；撤销重做保持基本可用。为了答辩演示，登录页加入 NEU 品牌元素，但业务逻辑仍以可用性优先。",
        "实现上我承认 app.js 体积偏大，这是技术债。短期为了赶联调把逻辑堆在单文件，长期应拆成病例模块、工具模块、AI 模块与手术模块。报告里把它写出来，是为了说明我清楚“能跑”和“好维护”不是一回事。",
    ]:
        add_para(doc, t)
    add_figure(doc, FIG / "fig_placeholder_ui.png", "图2-13 二维多标签标注（预留，待替换）")
    add_heading_cn(doc, "2.5.4  三维、手势与模拟手术", 3)
    for t in [
        "volume_viewer.js 管理 VTK/WebGL2 场景、MPR 联动与手术 overlay。表面网格在 per_label 模式下分层上色，避免多器官糊成一团。手势模块 hand_gesture.js 解析双手关键点，映射旋转、缩放、选器官等；进入前预检体数据层数与 TotalSeg 健康状态。",
        "模拟手术确认 ROI 后，投影框与 meshScale 对齐，修复“选中后绿框消失”的坐标不一致问题。保存时 getSurgerySnapshot() 带出器官字段，app.js 组装 POST /api/surgery_results。收刀后保留红色刀痕面，和长方体一起构成可解释的多面体 ROI。",
    ]:
        add_para(doc, t)
    add_figure(doc, FIG / "fig_placeholder_ui.png", "图2-14 三维视图 + 手势 dock + MPR（预留，待替换）")
    add_figure(doc, FIG / "fig_placeholder_ui.png", "图2-15 模拟手术 ROI 与刀痕面（预留，待替换）")
    add_heading_cn(doc, "2.5.5  AI 对接与协作发布", 3)
    for t in [
        "ai_service 调用预测时默认不允许 silent baseline；无模型或服务不可用返回明确错误。TotalSeg 对深度不足的体数据直接拒绝。平台暴露 model_status / backend / fallback_reason 等信息，方便现场解释当前后端。多器官 Plan A 接入后，前端模型列表需能选到对应后端，而不是写死单一器官。",
        "代码经 feature-a 推送，并已合并至远程 main 与 dev。合并时不把 .env、权重和大影像带进仓库；web/public 同步 volume_viewer.js / hand_gesture.js。发布前自检：强制刷新、用多层 CT 走三维与手术保存、确认库中写出器官名。",
    ]:
        add_para(doc, t)

    add_heading_cn(doc, "2.6  系统测试", 2)
    add_heading_cn(doc, "2.6.1  测试方案设计", 3)
    for t in [
        "测试目标：验证平台在真实运行环境下关键业务满足需求并可答辩演示。依据：docs/04 接口设计、docs/17_system_test_plan.md、docs/18_manual_ui_checklist.md、scripts/run_system_tests.sh。",
        "被测对象：本人负责的平台侧（FastAPI、SQLite、legacy frontend、手术 ROI、审核工作流）。测试类型：功能、权限/安全负向、边界异常、工作流集成、上传导出、性能冒烟、浏览器人工验收。",
        "环境：macOS，uvicorn @ 127.0.0.1:8000，SQLite，Chrome/Edge；演示数据 Case0002–0004；上传类用例使用 [SYSTEM_TEST] 独立病例。策略：等价类与边界；缺陷分阻断演示/影响体验/技术债。通过准则：阻断级全过；自动化失败为 0；人工 UI 全项通过。",
    ]:
        add_para(doc, t)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["被测版本", "feature-a / main 可运行代码"],
            ["自动化入口", "bash scripts/run_system_tests.sh"],
            ["报告产物", "docs/report/system_test_report.md"],
            ["人工检查", "docs/18_manual_ui_checklist.md"],
        ],
    )
    add_heading_cn(doc, "2.6.2  测试用例与结果", 3)
    add_para(
        doc,
        "自动化最新结果：共 56 项，通过 54，失败 0，跳过 2（长耗时 TotalSeg 实装预测与短训启动，可用 SYSTEM_TEST_RUN_HEAVY=1 开启），通过率 96.43%，结论 PASS。人工 UI 检查 15/15 全部通过（2026-07-14，王瑞琦）。",
    )
    add_figure(doc, FIG / "fig10_test_summary.png", "图2-16 系统测试结果汇总")
    add_table(
        doc,
        ["用例编号", "场景", "输入/步骤要点", "期望", "结果"],
        [
            ["ST-AUTH-01", "三角色登录", "annotator/reviewer/admin", "JWT 角色正确", "通过"],
            ["ST-AUTH-02", "越权写接口", "annotator 调审核接口", "拒绝（4xx）", "通过"],
            ["ST-CASE-01", "病例与切片", "打开多层 CT", "可出 PNG/体数据", "通过"],
            ["ST-MASK-01", "Mask 写路径", "save/update/promote", "版本可查可回滚", "通过"],
            ["ST-UP-01", "NIfTI 上传", "带 SYSTEM_TEST 标记", "新建 case/image", "通过"],
            ["ST-EX-01", "导出 materialize", "选定版本导出", "目录与 manifest", "通过"],
            ["ST-WF-01", "审核闭环", "submit→reject→resubmit→approve", "状态可追溯", "通过"],
            ["ST-SURG-01", "手术 ROI 入库", "确认 ROI 后保存", "含 organ_* 字段", "通过"],
            ["ST-SURG-02", "非法 ROI", "未确认/非法坐标", "拒绝入库", "通过"],
            ["ST-AI-01", "AI 诚实失败", "无模型/不可用", "明确错误原因", "通过"],
            ["ST-AI-H", "重 AI/短训", "SYSTEM_TEST_RUN_HEAVY", "可选执行", "跳过"],
            ["ST-PERF-01", "健康延迟冒烟", "反复 /health", "阈值内", "通过"],
            ["UI-VTK-01~04", "三维/MPR", "人工检查表", "渲染与联动正常", "通过"],
            ["UI-GES-01~04", "手势交互", "摄像头+MediaPipe", "映射与预检正确", "通过"],
            ["UI-SUR-01~05", "手术三步", "选器官→确认→切割", "门槛与入库正确", "通过"],
            ["UI-BR-01~02", "浏览器抽测", "Chrome/Edge", "行为一致", "通过"],
        ],
    )
    add_table(
        doc,
        ["分组", "条目数", "结果"],
        [
            ["自动化 HTTP 系统测试", "56（含 2 跳过）", "54 通过 / 0 失败"],
            ["人工 UI（VTK/手势/手术/浏览器）", "15", "15/15 通过"],
        ],
    )
    add_heading_cn(doc, "2.6.3  测试结论与评价", 3)
    for t in [
        "测试过程中已修复并回归的问题包括：ROI 投影与 meshScale 不一致导致绿框“消失”；握拳收刀识别不稳改为捏合；手势浮层遮挡改为文档流 dock；旧库缺器官列通过 ALTER 兼容；promote 到 final 前需满足 pending/reviewed 状态。",
        "残留风险：手势受光照影响；VTK/WebGL2 双路径维护成本；手术坐标未做临床级审计；长耗时 AI/训练在默认回归中跳过，答辩前如需展示应预热权重环境。",
        "最终评价：自动化系统测试 PASS，人工 UI 检查全部通过，主演示路径（标注—三维—手势—手术入库—审核导出）达到答辩可用标准。系统测试表明本人交付的平台能力在约定范围内可用、可复查；后续建议将重测试纳入夜间任务，并继续拆分前端模块以降低回归成本。",
    ]:
        add_para(doc, t)

    # ========== 3 ==========
    add_heading_cn(doc, "3  总结", 1)
    add_heading_cn(doc, "3.1  实训体会", 2)
    for t in [
        "最大的收获不是又记住某一个 API，而是第一次把接口契约、数据标准、可运行演示和分支合并当成同一件事。医学影像里很多缺陷出在坐标系、层厚和标签语义，表现却像“功能丢了”，必须回到空间变换与数据约定排查。",
        "和队友联调让我更清楚：平台既要接待模型成功，也要接待模型失败。以前写作业容易只演示成功路径；现在更在意失败路径和数据是否还能解释。这更接近软件工程专业真正要训练的能力。",
        "不足也很明显：前端单文件仍然偏大；长耗时 AI/训练默认未纳入每次回归；模拟手术距真实临床仍远。后续会把重测试纳入夜间任务，并继续拆分前端模块。",
        "从职业视角看，实训逼着我练习：把含糊需求落成可验收条目；在队友接口未就绪时用契约继续推进；在时间不够时做减法。这些比“又用了一个新框架”更重要。",
    ]:
        add_para(doc, t)
    add_heading_cn(doc, "3.2  其它意见", 2)
    for t in [
        "对学院：企业项目实训把文档、答辩与工程实践绑在一起，方向正确；若中期增加“接口冻结日”，组间联调会更顺。也建议提供标准样例包（层数、模态、已知问题）减少环境差异。",
        "对项目组织：把公开样例的层数与模态说明置顶到 README；尽量安排一次短代码走查，提前暴露“能跑但不可维护”的问题。对企业指导教师：封面姓名待补，亦希望能就工程边界与职业规范给出口头反馈。",
        "以上意见仅供参考。本人将继续补齐运行界面截图，并按评阅意见修改报告细节。",
    ]:
        add_para(doc, t)

    add_heading_cn(doc, "参考文献与附录说明", 1)
    for t in [
        "[1] 东北大学软件学院. 学生实训总结报告撰写说明（附件05）.",
        "[2] C 组项目文档 docs/01–16：数据标准、API 设计、ER、原型、GitHub 协作与 AI 联调说明.",
        "[3] FastAPI / VTK.js / MediaPipe / TotalSegmentator / nnU-Net 等相关技术文档（实训期间查阅）.",
        "[4] docs/17_system_test_plan.md、docs/18_manual_ui_checklist.md、docs/report/system_test_report.md（系统测试计划、人工检查表与测试报告）.",
        "附录：图2-3、图2-4、图2-12–图2-15 等为运行界面预留位，定稿前由本人替换为实际截图；数据库与接口以仓库当前代码为准。",
    ]:
        add_para(doc, t, first_line=False)
    add_para(doc, "（正文完）", first_line=False, align="center")

    text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            text += "\n" + "\t".join(c.text for c in row.cells)
    out = OUT_DIR / "王瑞琦-20236742-软英2301-学生实训总结报告-中文版.docx"
    doc.save(out)
    import re

    zh_n = len(re.findall(r"[\u4e00-\u9fff]", text))
    print("ZH saved:", out, "chinese_chars:", zh_n)
    return out, zh_n


def build_en():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.6)
    setup_page_header(
        sec,
        "School of Software, Northeastern University",
        "Student Internship Summary Report",
        en=True,
    )
    cover_en(doc)

    add_heading_en(doc, "Notes", 1)
    for t in [
        "Before the internship ends, every student must carefully write a Student Internship Summary Report, systematically reviewing the full process and raising practical experience to a theoretical level.",
        "The report must be clear, detailed, and accurate, with no fewer than 5,000 Chinese characters in the Chinese edition. Report quality is a major basis for grading; late or missing reports cannot be graded.",
        "The cover must use the full internship-base name and the industry mentor’s name (currently marked “[To be filled]”).",
        "This English edition mirrors the Chinese template structure: Preface → Internship Content (Overview, Related Technologies, System Analysis, System Design, System Implementation, System Testing) → Conclusion. It includes architecture, use-case, activity, class/object, state, sequence, collaboration, and program-flow diagrams, plus key code. UI figures marked as placeholders will be replaced with runtime screenshots.",
        "The surgery module is a demo-oriented ROI recorder, not a certified clinical surgical planner.",
    ]:
        add_para(doc, t)
    doc.add_page_break()

    add_heading_en(doc, "Contents", 1)
    for line in [
        "1  Preface",
        "    1.1  Internship Background",
        "    1.2  Internship Environment",
        "    1.3  Internship Process",
        "2  Internship Content",
        "    2.1  Overview",
        "    2.2  Related Technologies",
        "    2.3  System Analysis",
        "        2.3.1  Business Analysis",
        "        2.3.2  Functional Analysis (Use Cases & Activity Diagrams)",
        "        2.3.3  Feasibility and Scope",
        "    2.4  System Design",
        "        2.4.1  Overall Design (Architecture & Modules)",
        "        2.4.2  Detailed Design (Classes, Dynamics, Algorithms)",
        "        2.4.3  Database Design",
        "    2.5  System Implementation",
        "    2.6  System Testing",
        "        2.6.1  Test Plan Design",
        "        2.6.2  Test Cases and Results",
        "        2.6.3  Conclusion and Evaluation",
        "3  Conclusion",
        "    3.1  Reflections",
        "    3.2  Other Suggestions",
        "References and Appendix Notes",
    ]:
        add_para(doc, line, first_line=False, space_after=2)
    doc.add_page_break()

    add_heading_en(doc, "1  Preface", 1)
    add_heading_en(doc, "1.1  Internship Background", 2)
    for t in [
        "The enterprise project internship is placed in the senior undergraduate stage. The college expects more than a course-style page demo: students should push a medical-imaging-related system to a runnable, integrable, defense-ready prototype under team collaboration. Group C builds a CT-oriented annotation and data-management platform (repository: C_Group_Medical_Annotation).",
        "The domain tension is practical. Manual 3D delineation is slow and hard to version; fully automatic segmentation rarely passes review in one shot. The platform therefore targets a closed loop of AI draft labeling, human refinement, review, and dataset export. For a software-engineering student, this covers requirements, APIs, persistence, interaction, quality, and collaboration.",
        "We split platform engineering and AI algorithms. I owned the platform side: backend APIs, database, annotation frontend, 3D visualization, auth/review workflow, and the AI integration surface. My teammate focused on training/inference. We shared one GitHub repo, naming rules, and /api contracts. My goals were to finish a full SE lifecycle, confront messy DICOM/NIfTI/mask realities, and keep the system integrable when teammates moved at different speeds.",
        "Up front: this is an educational engineering system, not a certified medical device. “Simulated surgery” validates 3D interaction and ROI persistence; it is not clinical surgical planning.",
    ]:
        add_para(doc, t)

    add_heading_en(doc, "1.2  Internship Environment", 2)
    for t in [
        "The internship base is NEU Software College enterprise project training, carried by the Group C platform. Local stack: macOS; FastAPI/uvicorn on 127.0.0.1:8000; SQLite with schema.sql plus runtime ensure_*_schema migrations; primary UI in frontend/ (vanilla JS) with React assets under web/; VTK.js and WebGL2 volume rendering; MediaPipe Hands; AI bridges to nnU-Net, TotalSegmentator, and DeepEdit (local TotalSeg Python under conda when needed).",
        "GitHub branches: feature-a / feature-b / dev / main. Large images, weights, and real patient data stay out of the repository; only code, docs, configs, and split manifests are committed. Seed users (annotator / reviewer / admin) support permission debugging.",
        "Day-to-day work sat in the platform engineering lane: API design/implementation, workstation interaction, volume/mask services, gesture and surgery modules, teammate predict write-back integration, documentation, and branch merges. Please fill the industry mentor name on the cover where marked.",
    ]:
        add_para(doc, t)
    add_table(
        doc,
        ["Item", "Detail"],
        [
            ["OS", "macOS"],
            ["Backend", "FastAPI + uvicorn (:8000)"],
            ["Database", "SQLite + schema migration"],
            ["Frontend", "Vanilla workstation + React assets"],
            ["3D / Gestures", "VTK.js, WebGL2, MediaPipe"],
            ["AI", "nnU-Net / TotalSeg / DeepEdit / platform_unet"],
            ["Collaboration", "GitHub feature/dev/main branches"],
        ],
    )

    add_heading_en(doc, "1.3  Internship Process", 2)
    for t in [
        "Phase 1 built standards and skeleton: data naming, ER draft, API contract, UI wireframe, and a working cases/images/masks/versions path. Without that backbone, later 3D and AI features would have nowhere to land. The first milestone was intentionally boring: select a case, scroll slices, save a mask, reload it.",
        "Phase 2 added JWT auth, task assignment, submit/approve/reject, VTK 3D + MPR, MIP/MinIP, and multi-DICOM upload. The goal was to make “a day in an annotator’s life” possible inside one system instead of scattered scripts.",
        "Phase 3 deepened human–AI collaboration: honest predict failures, multi-label tools, quality/export, then MediaPipe bimanual navigation and organ picking, integrating my teammate’s multi-organ capabilities through platform APIs. This is also where we learned that a beautiful wrong mask is worse than a plain error toast.",
        "Phase 4 polished for defense: gesture dock between 3D and MPR, per-label colored VTK meshes, surgery gates (select → confirm cuboid ROI → cut), surgery_results with organ fields, and merging feature-a into main/dev on GitHub.",
        "Rework worth noting: configurable export manifests when my teammate’s schemas evolved; depth/health prechecks against single-slice demos; replacing an overlay gesture panel that blocked 3D with a document-flow dock. These changes are not glamorous, but they match how real products are finished under deadline pressure.",
        "Across phases I kept a simple personal rule: never claim a feature in the report unless the current branch can demonstrate it after a hard refresh. That rule prevented several “slides ahead of code” embarrassments during dry runs.",
    ]:
        add_para(doc, t)
    add_figure(doc, FIG_EN / "fig3_annotation_flow.png", "Fig. 1-1 Main annotation workflow")
    add_figure(doc, FIG_EN / "fig8_review_flow.png", "Fig. 1-2 Review workflow (simplified)")

    add_heading_en(doc, "2  Internship Content", 1)
    add_heading_en(doc, "2.1  Overview", 2)
    for t in [
        "The system serves annotators, reviewers, and admins for CT upload, 2D multi-label editing, 3D/gesture interaction, AI mask write-back, version review, dataset export, and demo surgery ROI recording. I did not train a brand-new SOTA network; I made the platform call models honestly and store results reproducibly.",
        "In product terms, I own the “operating system” of labeling: identity, persistence, visualization, workflow, and integration sockets. My teammate owns learning algorithms and training recipes. The internship grade, for my part, depends on whether those sockets are stable enough for another person to plug into.",
        "Deliverables spanned FastAPI/SQLite services, auth/task/review APIs, imaging/volume endpoints, mask versioning and export, the annotation workstation, VTK/WebGL2 + MPR, MediaPipe gestures, organ-aware surgery ROI persistence, teammate AI integration, and release merges to main/dev.",
        "The rest of Chapter 2 expands technologies, analysis, design, implementation, and testing in that order, matching the college template while staying anchored to code that exists in the repository today.",
    ]:
        add_para(doc, t)
    add_figure(doc, FIG_EN / "fig1_architecture.png", "Fig. 2-1 Overall architecture")
    add_figure(doc, FIG_EN / "fig7_modules.png", "Fig. 2-2 My module breakdown")
    add_figure(doc, FIG_EN / "fig_placeholder_ui.png", "Fig. 2-3 Login UI (placeholder)")
    add_figure(doc, FIG_EN / "fig_placeholder_ui.png", "Fig. 2-4 Workstation overview (placeholder)")

    add_heading_en(doc, "2.2  Related Technologies", 2)
    add_heading_en(doc, "2.2.1 Theoretical Background", 3)
    for t in [
        "Medical volumes appear as DICOM series or NIfTI/NRRD tensors; masks encode labels on the same grid. Multi-label scenes require careful label IDs so a liver edit does not erase a kidney. Dice/IoU support coarse version comparison but never replace expert review.",
        "Volume rendering and surface meshes (e.g., marching cubes) serve different reading goals. Volume rendering helps inspect density distributions; surfaces help understand organ boundaries and support picking. MPR remains essential for radiology-style correlation across axial, coronal, and sagittal views.",
        "Human–AI annotation assumes models propose and humans correct. Therefore the system must support AI write-back and human overwrite with version provenance. Gestures reduce mouse load in 3D but introduce recognition noise that must be gated by explicit modes and prechecks.",
        "From a software-architecture view, the hard problem is not drawing polygons; it is keeping identity stable across upload, prediction, edit, review, and export so that a later training run can explain where each label came from.",
    ]:
        add_para(doc, t)
    add_heading_en(doc, "2.2.2 Methods, Tools, and Environment", 3)
    for t in [
        "We used iterative increments: freeze contracts, ship a vertical slice, then deepen 3D/gestures/surgery under defense pressure. Docs 01–16 acted as team manuals and reduced lost oral agreements.",
        "Tooling included FastAPI/Pydantic, vanilla JS plus React assets, VTK.js/WebGL2, MediaPipe, and Python AI stacks. Local uvicorn supported fast backend iteration; node --check and py_compile acted as cheap syntax gates before demos.",
        "Methodologically we avoided pretending to run a heavyweight waterfall. When requirements changed—for example “surgery must confirm ROI first”—we patched the running system in small steps and updated the contract docs in the same pull request whenever possible.",
    ]:
        add_para(doc, t)
    add_table(
        doc,
        ["Category", "Choice", "Rationale"],
        [
            ["Web API", "FastAPI", "Typed contracts, fast OpenAPI"],
            ["DB", "SQLite", "Light deploy for training scale"],
            ["UI", "Vanilla JS workstation", "Few demo dependencies"],
            ["3D", "VTK.js + WebGL2", "Surface/volume options"],
            ["Gestures", "MediaPipe Hands", "Browser-ready"],
            ["AI", "nnU-Net/TotalSeg/DeepEdit", "Aligned with teammate AI"],
        ],
    )

    add_heading_en(doc, "2.3  System Analysis", 2)
    add_para(
        doc,
        "System analysis moves from “what problem the whole project solves” to “what I must deliver.” The subsections cover business, functions (use cases and activity diagrams), and feasibility.",
    )
    add_heading_en(doc, "2.3.1  Business Analysis", 3)
    for t in [
        "A case flows upload → browse → annotate (optionally AI-first) → save mask/version → submit → review → export. Multi-organ labels, coexisting versions, recoverable failures, and role isolation matter throughout that chain.",
        "For 3D scenes, users also want rotation, MPR correlation, optional low-mouse gestures, and a demo “pseudo-surgery” ROI on a selected organ. These needs do not replace clinical surgical planning, but they test whether interaction and persistence actually close.",
        "If the loop breaks at “can view but cannot save,” or “saved but cannot tell which organ was cut,” the system’s value collapses quickly in a defense setting. That is why organ metadata was promoted from a UI-only concern into first-class database fields.",
        "Stakeholders differ: annotators want speed and undo; reviewers want diffs and accountability; admins want assignment; AI teammates want stable export schemas. My platform has to answer all of them, not only polish one animation.",
    ]:
        add_para(doc, t)
    add_heading_en(doc, "2.3.2  Functional Analysis (Use Cases & Activity Diagrams)", 3)
    add_para(doc, "Figure 2-5 summarizes actors and use cases. Authorization is enforced in APIs, not only by hiding buttons.")
    add_figure(doc, FIG_EN / "fig2_usecase.png", "Fig. 2-5 Core use cases")
    add_table(
        doc,
        ["Use case", "Actor", "Precondition", "Postcondition"],
        [
            ["Login", "All", "Account exists", "JWT + role"],
            ["Upload CT", "Annotator/Admin", "Logged in", "case/image created"],
            ["Save annotation", "Annotator", "Image open", "mask/version stored"],
            ["AI predict", "Annotator", "Model available", "AI mask or honest error"],
            ["3D/gestures", "Annotator", "Enough slices", "Navigate/pick organ"],
            ["Surgery ROI", "Annotator", "Organ + confirmed ROI", "surgery_results row"],
            ["Review", "Reviewer", "Submitted", "Approved/rejected"],
            ["Export dataset", "Annotator/Admin", "Usable version", "Train tree + manifest"],
        ],
    )
    add_para(
        doc,
        "Activity diagrams capture cross-role processes. Fig. 2-5a shows the main annotation activity (upload to export); Fig. 2-5b shows review (submit—approve/reject—re-edit). They align with later state and sequence diagrams.",
    )
    add_figure(doc, FIG_EN / "fig3_annotation_flow.png", "Fig. 2-5a Annotation activity diagram")
    add_figure(doc, FIG_EN / "fig8_review_flow.png", "Fig. 2-5b Review activity diagram")
    add_heading_en(doc, "2.3.3  Feasibility and Scope", 3)
    for t in [
        "Technically and economically feasible with open-source stacks and local SQLite. Operationally, browser clients cover login/annotation/3D/gestures; AI runs on the agreed local Python environment. Ethically, only teaching/public sample volumes are used—no real patient PHI.",
        "Scope honesty: teaching-oriented 2.5D platform U-Net; high-accuracy paths still rely on TotalSeg/nnU-Net; surgery is interactive demonstration; gestures are lighting-sensitive. Single-slice volumes are rejected for 3D demos.",
    ]:
        add_para(doc, t)
    add_table(
        doc,
        ["Priority", "Capability", "Note"],
        [
            ["P0", "Upload/browse/annotate/review/export", "Core labeling system"],
            ["P0", "Honest AI predict", "No silent fake masks"],
            ["P1", "VTK 3D / MPR / MIP-MinIP", "Reading and QC"],
            ["P1", "Gestures + organ focus", "Lower 3D friction"],
            ["P2", "Surgery ROI persistence", "Demo closed loop"],
        ],
    )

    add_heading_en(doc, "2.4  System Design", 2)
    add_heading_en(doc, "2.4.1  Overall Design (Architecture & Modules)", 3)
    for t in [
        "Layering is presentation → API → services → data/files (Fig. 2-1). The frontend never talks to SQLite directly; AI does not silently overwrite formal versions. That separation keeps audit and rollback possible when a prediction is wrong.",
        "Modules follow Figure 2-2: auth/review, imaging, masks/versions, AI proxy, 3D rendering, gestures, surgery ROI, and the workstation shell. I own backend/frontend/database/docs; my teammate owns ai/; both share dataset conventions under dataset/.",
        "Cross-cutting concerns include JWT authz, consistent ID naming (Case0001/Image0001/Mask0001), and schema migration helpers so an old demo database does not explode on the defense machine.",
    ]:
        add_para(doc, t)
    add_heading_en(doc, "2.4.2  Detailed Design (Classes, Dynamics, Algorithms)", 3)
    add_para(
        doc,
        "Detailed design centers on my work: class/object structure, dynamic models, and algorithms. The backend is organized as route schemas + service functions + domain records rather than a heavy ORM hierarchy; the frontend uses collaborating module objects.",
    )
    add_figure(doc, FIG_EN / "fig11_class.png", "Fig. 2-6a Core classes / objects")
    add_para(
        doc,
        "Dynamic models include a state diagram (review + surgery gates, Fig. 2-6b), a sequence diagram (save surgery ROI, Fig. 2-8), and a collaboration diagram (AI predict write-back, Fig. 2-6c).",
    )
    add_figure(doc, FIG_EN / "fig12_state.png", "Fig. 2-6b State machines")
    add_figure(doc, FIG_EN / "fig6_sequence.png", "Fig. 2-8 Save-surgery sequence")
    add_figure(doc, FIG_EN / "fig14_collab.png", "Fig. 2-6c AI predict collaboration")
    add_figure(doc, FIG_EN / "fig4_surgery_flow.png", "Fig. 2-7 Surgery ROI steps")
    add_figure(doc, FIG_EN / "fig9_gesture_map.png", "Fig. 2-9 Gesture-to-action mapping")
    for t in [
        "Algorithm design: (1) organ-field resolution—request organ object first, then label catalog fallback; (2) ROI validity—cuboid_max≥cuboid_min, label_id>0, case/image ownership; (3) per_label surface extraction for colored multi-organ meshes; (4) Dice/IoU compare for coarse version disagreement, never replacing human review.",
        "APIs use the /api prefix, including auth, cases, upload, slices/volume, masks, predict, export, and surgery_results. Surface meshes support per_label marching cubes for colored VTK layers.",
    ]:
        add_para(doc, t)
    add_heading_en(doc, "2.4.3  Database Design", 3)
    for t in [
        "Core entities include users, cases, images, annotations, masks, versions, models, tasks, and surgery_results. cases is the hub; images belong to cases; masks attach labels to images; versions freeze a reviewable state; surgery_results records demo ROIs.",
        "Surgery rows store cuboid bounds, cut planes, carved voxels, plus organ_name / organ_display_name / organ_color / organ_json so a saved ROI can answer which organ was targeted. ensure_surgery_schema migrates older DBs with PRAGMA table_info + ALTER TABLE instead of forcing a wipe.",
    ]:
        add_para(doc, t)
    add_figure(doc, FIG_EN / "fig5_er.png", "Fig. 2-10 Selected entity relationships")
    add_table(
        doc,
        ["Table", "Key fields", "Role"],
        [
            ["cases", "case_id, status", "Case hub"],
            ["images", "image_id, path, shape", "Volume metadata"],
            ["masks", "mask_id, label_id, path", "Segmentation store"],
            ["versions", "version_id, tag", "Provenance"],
            ["users/tasks", "role, assignee", "Authz and work"],
            ["surgery_results", "organ_*, cuboid_*, planes", "ROI + organ info"],
        ],
    )

    add_heading_en(doc, "2.5  System Implementation", 2)
    add_para(
        doc,
        "Implementation followed “vertical slice first, then edges.” The subsections below expand my deliverables with a program flowchart, key code, and UI placeholders.",
    )
    add_heading_en(doc, "2.5.1  Backend, Auth, and Persistence", 3)
    for t in [
        "main.py registers routers; services own business rules; SQLite access is centralized. JWT login returns a bearer token; protected write endpoints check roles instead of trusting hidden buttons. Task and review APIs move case state and append audit information so “who approved what” remains answerable.",
        "Saving surgery results validates case/image ownership and label_id, normalizes cut planes, and resolves organ metadata from the request body or label catalog before INSERT. Missing display names fall back to catalog values so the database is not left with a bare integer.",
    ]:
        add_para(doc, t)
    add_figure(doc, FIG_EN / "fig13_program_flow.png", "Fig. 2-11 Program flow: save surgery ROI")
    add_code(
        doc,
        "POST /api/surgery_results\n"
        "body: case_id, image_id, mask_id?, label_id, organ_*,\n"
        "      roi_margin_pct, knife_radius, cuboid_min/max,\n"
        "      cut_planes[], carved_voxels, note?",
    )
    add_code(
        doc,
        "# surgery_service.py (semantic excerpt)\n"
        "def save_surgery_result(request, user=None):\n"
        "    ensure_surgery_schema()\n"
        "    case = get_record('cases', 'case_id', request.case_id)\n"
        "    if case is None: raise HTTPException(404, ...)\n"
        "    image = get_record('images', 'image_id', request.image_id)\n"
        "    if image['case_id'] != request.case_id: raise HTTPException(422, ...)\n"
        "    if request.label_id <= 0: raise HTTPException(422, ...)\n"
        "    organ = _resolve_organ_fields(request)  # request first, catalog fallback\n"
        "    # INSERT surgery_results(... organ_name/display/color ...)\n"
        "    return SaveSurgeryResultResponse(...)",
    )
    add_table(
        doc,
        ["Field", "Meaning", "Source"],
        [
            ["label_id", "Organ label id", "Current selection"],
            ["organ_name", "Internal name", "Request or catalog"],
            ["organ_display_name", "UI name", "Request or catalog"],
            ["organ_color", "Display color", "Request or catalog"],
            ["cuboid_min/max", "ROI corners", "Confirmed cuboid"],
            ["cut_planes", "Cut face list", "Accumulated during cutting"],
            ["carved_voxels", "Rough carved count", "Frontend statistic"],
        ],
    )
    add_heading_en(doc, "2.5.2  Imaging, Masks, and Export", 3)
    for t in [
        "Upload handles NIfTI packs and multi-DICOM series; the backend stores files, creates image rows, and serves slice PNGs plus volume metadata.",
        "Masks are label-aware; versions separate manual/ai/final semantics; compare exposes Dice/IoU for quick disagreement checks. Export can materialize multiclass labels plus dataset.json/manifests so my teammate can train against a stable tree.",
    ]:
        add_para(doc, t)
    add_figure(doc, FIG_EN / "fig_placeholder_ui.png", "Fig. 2-12 Version/compare panel (placeholder)")
    add_heading_en(doc, "2.5.3  Annotation Workstation", 3)
    for t in [
        "app.js orchestrates case list, tools, AI actions, versions, save, and export. Multi-label brush/eraser honor the active label_id; undo/redo stay basic but usable.",
        "I admit app.js is oversized. Short-term integration pressure pushed logic into one file; longer-term it should split into case, tool, AI, and surgery modules.",
    ]:
        add_para(doc, t)
    add_figure(doc, FIG_EN / "fig_placeholder_ui.png", "Fig. 2-13 2D multi-label annotation (placeholder)")
    add_heading_en(doc, "2.5.4  3D, Gestures, and Surgery", 3)
    for t in [
        "volume_viewer.js owns VTK/WebGL2 scenes, MPR sync, and surgery overlays. In per_label mode, surfaces are colored by organ so multi-class scenes remain readable.",
        "hand_gesture.js maps landmarks to orbit/zoom/pick. Before opening the camera it checks volume depth and TotalSeg health. After ROI confirmation, projection boxes align with meshScale. getSurgerySnapshot() exports organ fields; app.js posts them to /api/surgery_results.",
    ]:
        add_para(doc, t)
    add_figure(doc, FIG_EN / "fig_placeholder_ui.png", "Fig. 2-14 3D + gesture dock + MPR (placeholder)")
    add_figure(doc, FIG_EN / "fig_placeholder_ui.png", "Fig. 2-15 Surgery ROI and cut faces (placeholder)")
    add_heading_en(doc, "2.5.5  AI Integration and Release", 3)
    for t in [
        "ai_service rejects silent baselines; shallow volumes are blocked on TotalSeg paths; model_status/backend/fallback_reason help explain live demos.",
        "feature-a was merged to main/dev without secrets or giant weights. Pre-release self-check: hard refresh, open a multi-slice CT, walk surgery save, confirm organ names in SQLite.",
    ]:
        add_para(doc, t)

    add_heading_en(doc, "2.6  System Testing", 2)
    add_heading_en(doc, "2.6.1  Test Plan Design", 3)
    for t in [
        "Goal: verify that the platform meets requirements in a real runtime and is defense-ready. Basis: docs/04, docs/17_system_test_plan.md, docs/18_manual_ui_checklist.md, scripts/run_system_tests.sh.",
        "SUT: the platform modules I own (FastAPI, SQLite, legacy frontend, surgery ROI, review workflow). Types: functional, auth/security negatives, boundary, workflow integration, upload/export, performance smoke, browser manual acceptance.",
        "Environment: macOS, uvicorn @ 127.0.0.1:8000, SQLite, Chrome/Edge; demo volumes Case0002–0004; [SYSTEM_TEST] patients for uploads. Pass criteria: zero automated failures and all manual UI items passed.",
    ]:
        add_para(doc, t)
    add_table(
        doc,
        ["Item", "Detail"],
        [
            ["Build under test", "Runnable feature-a / main"],
            ["Automation entry", "bash scripts/run_system_tests.sh"],
            ["Report artifact", "docs/report/system_test_report.md"],
            ["Manual checklist", "docs/18_manual_ui_checklist.md"],
        ],
    )
    add_heading_en(doc, "2.6.2  Test Cases and Results", 3)
    add_para(
        doc,
        "Latest automated run: 56 cases, 54 passed, 0 failed, 2 skipped (heavy TotalSeg / short train; enable with SYSTEM_TEST_RUN_HEAVY=1), pass rate 96.43%, verdict PASS. Manual UI checklist: 15/15 passed (2026-07-14, Wang Ruiqi).",
    )
    add_figure(doc, FIG_EN / "fig10_test_summary.png", "Fig. 2-16 System test summary")
    add_table(
        doc,
        ["Case ID", "Scenario", "Key steps / inputs", "Expectation", "Result"],
        [
            ["ST-AUTH-01", "Tri-role login", "annotator/reviewer/admin", "Correct JWT roles", "Pass"],
            ["ST-AUTH-02", "Escalation write", "annotator hits review API", "Reject 4xx", "Pass"],
            ["ST-CASE-01", "Case & slices", "Open multi-slice CT", "PNG/volume OK", "Pass"],
            ["ST-MASK-01", "Mask write path", "save/update/promote", "Versions rollback OK", "Pass"],
            ["ST-UP-01", "NIfTI upload", "SYSTEM_TEST tag", "case/image created", "Pass"],
            ["ST-EX-01", "Export materialize", "Chosen version", "Tree + manifest", "Pass"],
            ["ST-WF-01", "Review loop", "submit→reject→resubmit→approve", "Traceable states", "Pass"],
            ["ST-SURG-01", "Surgery ROI save", "Confirm ROI then save", "organ_* persisted", "Pass"],
            ["ST-SURG-02", "Illegal ROI", "Unconfirmed / bad coords", "Reject save", "Pass"],
            ["ST-AI-01", "Honest AI failure", "No model / unavailable", "Explicit error", "Pass"],
            ["ST-AI-H", "Heavy AI/train", "SYSTEM_TEST_RUN_HEAVY", "Optional", "Skip"],
            ["ST-PERF-01", "Health latency", "Repeated /health", "Within threshold", "Pass"],
            ["UI-VTK-01~04", "3D / MPR", "Manual checklist", "Render & sync OK", "Pass"],
            ["UI-GES-01~04", "Gestures", "Camera + MediaPipe", "Mapping & precheck OK", "Pass"],
            ["UI-SUR-01~05", "Surgery steps", "select→confirm→cut", "Gates & save OK", "Pass"],
            ["UI-BR-01~02", "Browsers", "Chrome/Edge", "Consistent", "Pass"],
        ],
    )
    add_table(
        doc,
        ["Group", "Count", "Result"],
        [
            ["Automated HTTP system tests", "56 (2 skipped)", "54 pass / 0 fail"],
            ["Manual UI (VTK/gesture/surgery/browser)", "15", "15/15 pass"],
        ],
    )
    add_heading_en(doc, "2.6.3  Conclusion and Evaluation", 3)
    for t in [
        "Issues found and fixed include ROI projection vs meshScale mismatch, unstable fist-close replaced by pinch, overlay dock rewritten into document flow, ALTER migration for organ columns, and enforcing pending/reviewed before promote-to-final.",
        "Residual risks: lighting-sensitive gestures; dual VTK/WebGL2 maintenance; non-clinical surgery coordinates; heavy AI/train skipped in the default regression.",
        "Final evaluation: automated system testing PASS and manual UI checks fully passed. The defense path (annotate → 3D → gesture → surgery save → review/export) is usable and reviewable within my agreed scope.",
    ]:
        add_para(doc, t)

    add_heading_en(doc, "3  Conclusion", 1)
    add_heading_en(doc, "3.1  Reflections", 2)
    for t in [
        "The main takeaway is treating contracts, data standards, runnable demos, and merges as one job. Many medical-imaging defects are spatial/semantic rather than syntactic—when a green ROI box “vanishes,” the bug is often a transform, not a missing button handler.",
        "Working with my teammate reinforced that platforms must host both model success and model failure. Earlier coursework tempted me to demo only happy paths; this internship pushed failure paths and data explainability into the definition of done.",
        "Weaknesses remain: oversized frontend files, heavy AI/train still optional in default regression, and demo-grade surgery. If I continue, I will schedule nightlies for heavy jobs, modularize the UI, and add a surgery-result browser for live demos.",
        "Professionally, the internship trained requirement distillation, contract-first progress under incomplete dependencies, and deliberate scope cuts. Those habits matter more than collecting another framework name for a resume line.",
    ]:
        add_para(doc, t)
    add_heading_en(doc, "3.2  Other Suggestions", 2)
    for t in [
        "For the college: an mid-term API freeze day and a standard sample pack (modality, slice count, known pitfalls) would reduce integration thrash across machines.",
        "For project ops: put volume-depth constraints atop README and schedule a short code walkthrough to catch “runs but unmaintainable” early. Industry mentor name remains to be filled on the cover; spoken feedback on engineering boundaries would also help.",
        "I will keep replacing UI placeholders with real screenshots and revise details according to reviewer comments.",
    ]:
        add_para(doc, t)

    add_heading_en(doc, "References and Appendix Notes", 1)
    for t in [
        "[1] NEU Software College. Instructions for the Student Internship Summary Report (Attachment 05).",
        "[2] Group C docs/01–16: data standards, API design, ER, prototype, GitHub workflow, AI joint-debug notes.",
        "[3] FastAPI / VTK.js / MediaPipe / TotalSegmentator / nnU-Net documentation consulted during the internship.",
        "[4] docs/17_system_test_plan.md, docs/18_manual_ui_checklist.md, docs/report/system_test_report.md (system test plan, manual UI checklist, and test report).",
        "Appendix A: Figs. 2-3, 2-4, 2-12–2-15 are UI placeholders to be replaced with real screenshots; schema and APIs follow the current repository.",
        "Appendix B (defense checklist): use multi-slice cases; hard-refresh the browser; verify JWT role; run surgery save once and confirm organ_display_name in SQLite; keep .env and weights offline.",
        "Appendix C (limitations restated): platform U-Net is teaching-oriented; surgery ROI is not clinical planning; gesture quality depends on camera and lighting; heavy AI/train are optional in default regression.",
    ]:
        add_para(doc, t, first_line=False)
    add_figure(doc, FIG_EN / "fig_placeholder_ui.png", "Fig. A-1 Extra UI slot for final screenshot montage (placeholder)")
    add_para(doc, "(End of report)", first_line=False, align="center")

    text = "\n".join(p.text for p in doc.paragraphs)
    out = OUT_DIR / "WangRuiqi-20236742-Internship-Summary-Report-EN.docx"
    doc.save(out)
    print("EN saved:", out, "chars:", len(text.replace(" ", "").replace("\n", "")))
    return out, len(text.replace(" ", "").replace("\n", ""))


def main():
    zh, zh_n = build_zh()
    en, en_n = build_en()
    zh_alias = OUT_DIR / "王瑞琦-20236742-学生实训总结报告-中文翻译版.docx"
    en_alias = OUT_DIR / "王瑞琦-20236742-学生实训总结报告-英文版.docx"
    zh_alias.write_bytes(zh.read_bytes())
    en_alias.write_bytes(en.read_bytes())
    print("aliases updated")
    print("DONE chinese_chars=", zh_n, "en_chars~", en_n)


if __name__ == "__main__":
    main()
