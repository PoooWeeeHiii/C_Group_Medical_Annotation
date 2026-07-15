#!/usr/bin/env python3
"""
Fill the official NEU 附件05 template with internship report content,
preserving original paragraph/run formatting, section structure, and page setup.
"""

from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent
FIG = ROOT / "figures"
FIG_EN = ROOT / "figures" / "en"
OFFICIAL_DOC = Path(
    "/Users/wangruiqi/Desktop/实训/请仔细查看压缩包内  2023级企业项目实训结题答辩所用材料及说明"
    "/请仔细查看压缩包内  2023级企业项目实训结题答辩所用材料及说明"
    "/附件05：东北大学软件学院学生实训总结报告.doc"
)


def convert_official_doc() -> Path:
    """Convert official .doc -> .docx via textutil (keeps fonts better than pandoc)."""
    src_copy = ROOT / "附件05_源模板.doc"
    out = ROOT / "_official_base.docx"
    shutil.copy2(OFFICIAL_DOC, src_copy)
    import subprocess

    subprocess.run(
        ["textutil", "-convert", "docx", "-output", str(out), str(src_copy)],
        check=True,
    )
    return out


def set_paragraph_bottom_border(paragraph, sz="12", color="000000"):
    pPr = paragraph._p.get_or_add_pPr()
    # remove old border if any
    for child in list(pPr):
        if child.tag == qn("w:pBdr"):
            pPr.remove(child)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def clear_paragraph_keep_fmt(paragraph):
    """Remove all runs but keep pPr."""
    for child in list(paragraph._p):
        if child.tag == qn("w:r"):
            paragraph._p.remove(child)


def set_paragraph_text_keep_fmt(paragraph, text: str, *, prefer_run_idx: int | None = None):
    """Replace paragraph text while cloning rPr from an existing run."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    src = runs[prefer_run_idx] if prefer_run_idx is not None and prefer_run_idx < len(runs) else None
    if src is None:
        # pick first non-empty run, else first
        src = next((r for r in runs if r.text.strip()), runs[0])
    rPr = deepcopy(src._element.rPr) if src._element.rPr is not None else None
    clear_paragraph_keep_fmt(paragraph)
    new_r = OxmlElement("w:r")
    if rPr is not None:
        new_r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_r.append(t)
    paragraph._p.append(new_r)


def clone_body_paragraph_after(anchor, text: str, fmt_src):
    """Insert a new paragraph after anchor, copying pPr/rPr from fmt_src."""
    new_p = OxmlElement("w:p")
    # copy pPr
    src_pPr = fmt_src._p.find(qn("w:pPr"))
    if src_pPr is not None:
        new_p.append(deepcopy(src_pPr))
    # pick rPr from fmt_src
    rPr = None
    for r in fmt_src.runs:
        if r._element.rPr is not None:
            rPr = deepcopy(r._element.rPr)
            break
    new_r = OxmlElement("w:r")
    if rPr is not None:
        new_r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_r.append(t)
    new_p.append(new_r)
    anchor._p.addnext(new_p)
    # return a proxy-like element; caller may chain by finding next
    return new_p


def insert_picture_after(anchor, image_path: Path, width_cm: float = 14.0):
    """Insert a centered picture paragraph after anchor; returns the new paragraph element."""
    from docx.oxml.ns import nsmap
    from docx.text.paragraph import Paragraph

    # Create empty paragraph then add picture via python-docx API on a temporary approach:
    new_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    new_p.append(pPr)
    anchor._p.addnext(new_p)
    # Wrap as Paragraph to use add_run().add_picture
    para = Paragraph(new_p, anchor._parent)
    run = para.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return new_p


def insert_caption_after(anchor_el, text: str, fmt_src):
    new_p = OxmlElement("w:p")
    src_pPr = fmt_src._p.find(qn("w:pPr"))
    if src_pPr is not None:
        pPr = deepcopy(src_pPr)
        # force center
        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), "center")
        new_p.append(pPr)
    else:
        pPr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)
        new_p.append(pPr)
    rPr = None
    for r in fmt_src.runs:
        if r._element.rPr is not None:
            rPr = deepcopy(r._element.rPr)
            break
    new_r = OxmlElement("w:r")
    if rPr is not None:
        new_r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_r.append(t)
    new_p.append(new_r)
    anchor_el.addnext(new_p)
    return new_p


def find_para(doc, pred):
    for i, p in enumerate(doc.paragraphs):
        if pred(p.text):
            return i, p
    raise KeyError("paragraph not found")


def normalize_heading(text: str) -> str:
    t = text.strip()
    t = t.replace("\t", " ")
    t = re.sub(r"\s+", " ", t)
    return t


def setup_header_from_template(doc, left: str, right: str):
    """Create real page header; cover (first page) empty; remove trailing body header-like para."""
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    # clear first page header
    fp = sec.first_page_header
    for p in fp.paragraphs:
        clear_paragraph_keep_fmt(p)
        p.text = ""
    # body header
    header = sec.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    clear_paragraph_keep_fmt(hp)
    # copy rPr from trailing header-like body para if present
    trail = None
    for p in doc.paragraphs:
        if "东北大学软件学院" in p.text and "学生实训总结报告" in p.text and len(p.text) > 20:
            trail = p
            break
    rPr = None
    if trail is not None:
        for r in trail.runs:
            if r._element.rPr is not None:
                rPr = deepcopy(r._element.rPr)
                break
    usable = sec.page_width - sec.left_margin - sec.right_margin
    hp.paragraph_format.tab_stops.add_tab_stop(usable, WD_TAB_ALIGNMENT.RIGHT)
    # left run
    r1 = hp.add_run(left)
    if rPr is not None and r1._element.rPr is None:
        r1._element.insert(0, deepcopy(rPr))
    hp.add_run("\t")
    r2 = hp.add_run(right)
    if rPr is not None and r2._element.rPr is None:
        r2._element.insert(0, deepcopy(rPr))
    set_paragraph_bottom_border(hp, sz="12")
    sec.header_distance = Cm(1.5)
    # remove trailing body header duplicate(s)
    for p in list(doc.paragraphs):
        if "东北大学软件学院" in p.text and "学生实训总结报告" in p.text and len(p.text) > 20:
            # only remove if not the cover title pair (cover titles are short separate paras)
            if p.text.strip().startswith("东北大学软件学院") and "学生实训总结报告" in p.text:
                # if it's the combined header line, remove
                if len(p.text.strip()) > 25:
                    p._element.getparent().remove(p._element)


def fill_cover_zh(doc):
    rules = [
        (lambda t: t.startswith("专") and "业" in t[:8], "专    业： 软件工程（软件英才班）"),
        (lambda t: t.startswith("班") and "级" in t[:8], "班    级： 软英2301"),
        (lambda t: t.startswith("学") and "号" in t[:8], "学    号： 20236742"),
        (lambda t: t.startswith("姓") and "名" in t[:8], "姓    名： 王瑞琦"),
        (lambda t: t.startswith("实训基地"), "实训基地： 东北大学软件学院企业项目实训"),
        (lambda t: t.startswith("企业指导教师"), "企业指导教师： 【请填写】"),
    ]
    for p in doc.paragraphs[:25]:
        t = p.text.strip()
        for pred, val in rules:
            if pred(t):
                set_paragraph_text_keep_fmt(p, val)
                break


def fill_cover_en(doc):
    rules = [
        (lambda t: t.startswith("专") and "业" in t[:8], "专    业： Software Engineering (Elite Class)"),
        (lambda t: t.startswith("班") and "级" in t[:8], "班    级： SE-EN 2301"),
        (lambda t: t.startswith("学") and "号" in t[:8], "学    号： 20236742"),
        (lambda t: t.startswith("姓") and "名" in t[:8], "姓    名： Wang Ruiqi"),
        (lambda t: t.startswith("实训基地"), "实训基地： NEU Software College Enterprise Project Training"),
        (lambda t: t.startswith("企业指导教师"), "企业指导教师： [To be filled]"),
    ]
    for p in doc.paragraphs[:25]:
        t = p.text.strip()
        for pred, val in rules:
            if pred(t):
                set_paragraph_text_keep_fmt(p, val)
                break


def extract_sections_from_generated(path: Path, lang: str) -> dict[str, list]:
    """
    Extract body blocks from previously generated report.
    Returns dict: section_key -> list of items ('text', str) or ('fig', path, caption)
    """
    doc = Document(str(path))
    paras = [p.text.strip() for p in doc.paragraphs]

    def find_idx(keys, start_at=0):
        # Prefer body headings: skip TOC-indented lines from original raw text
        raw = [p.text for p in doc.paragraphs]
        for i in range(start_at, len(paras)):
            # TOC lines in our generated doc often start with 4 spaces
            if raw[i].startswith("    "):
                continue
            t = paras[i]
            if any(t == k or t.startswith(k + " ") or t.startswith(k + "\t") or t.startswith(k) for k in keys):
                # exact-ish: heading should be short
                if len(t) > 80:
                    continue
                return i
        return None

    def slice_between(start_keys, end_keys, start_at=0):
        start = find_idx(start_keys, start_at)
        if start is None:
            return []
        end = find_idx(end_keys, start + 1)
        if end is None:
            end = len(paras)
        out = []
        for t in paras[start + 1 : end]:
            if not t:
                continue
            if t.startswith(("图", "Fig.", "表", "Table")) and len(t) < 80:
                continue
            if t.startswith("[Missing"):
                continue
            out.append(("text", t))
        return out

    # Body starts at first non-TOC "1  前言" / "1  Preface"
    body_at = find_idx(["1  前言", "1  Preface"], 0) or 0

    if lang == "zh":
        sections = {
            "1.1": slice_between(["1.1  实训背景"], ["1.2  实训环境"], body_at),
            "1.2": slice_between(["1.2  实训环境"], ["1.3  实训过程"], body_at),
            "1.3": slice_between(["1.3  实训过程"], ["2  实训内容"], body_at),
            "2.1": slice_between(["2.1  概述"], ["2.2  相关技术"], body_at),
            "2.2": slice_between(["2.2  相关技术"], ["2.3  系统分析"], body_at),
            "2.3": slice_between(["2.3  系统分析"], ["2.4  系统设计"], body_at),
            "2.4": slice_between(["2.4  系统设计"], ["2.5  系统实现"], body_at),
            "2.5": slice_between(["2.5  系统实现"], ["2.6  系统测试"], body_at),
            "2.6": slice_between(["2.6  系统测试"], ["3  总结"], body_at),
            "3.1": slice_between(["3.1  实训体会"], ["3.2  其它意见"], body_at),
            "3.2": slice_between(["3.2  其它意见"], ["参考文献与附录说明", "参考文献"], body_at),
        }
        fig_plan = {
            "1.3": [
                (FIG / "fig3_annotation_flow.png", "图1-1 标注主流程"),
                (FIG / "fig8_review_flow.png", "图1-2 审核工作流活动图"),
            ],
            "2.1": [
                (FIG / "fig1_architecture.png", "图2-1 系统总体架构"),
                (FIG / "fig7_modules.png", "图2-2 本人负责的功能模块分解"),
            ],
            "2.3": [
                (FIG / "fig2_usecase.png", "图2-5 核心用例图"),
            ],
            "2.4": [
                (FIG / "fig11_class.png", "图2-6a 核心类与对象设计"),
                (FIG / "fig12_state.png", "图2-6b 状态图"),
                (FIG / "fig6_sequence.png", "图2-8 保存手术 ROI 时序图"),
                (FIG / "fig14_collab.png", "图2-6c AI 预测写回协作图"),
                (FIG / "fig5_er.png", "图2-10 核心数据实体关系"),
            ],
            "2.5": [
                (FIG / "fig13_program_flow.png", "图2-11 保存手术 ROI 程序流程图"),
            ],
            "2.6": [
                (FIG / "fig10_test_summary.png", "图2-16 系统测试结果汇总"),
            ],
        }
    else:
        sections = {
            "1.1": slice_between(["1.1  Internship Background"], ["1.2  Internship Environment"], body_at),
            "1.2": slice_between(["1.2  Internship Environment"], ["1.3  Internship Process"], body_at),
            "1.3": slice_between(["1.3  Internship Process"], ["2  Internship Content"], body_at),
            "2.1": slice_between(["2.1  Overview"], ["2.2  Related Technologies"], body_at),
            "2.2": slice_between(["2.2  Related Technologies"], ["2.3  System Analysis"], body_at),
            "2.3": slice_between(["2.3  System Analysis"], ["2.4  System Design"], body_at),
            "2.4": slice_between(["2.4  System Design"], ["2.5  System Implementation"], body_at),
            "2.5": slice_between(["2.5  System Implementation"], ["2.6  System Testing"], body_at),
            "2.6": slice_between(["2.6  System Testing"], ["3  Conclusion"], body_at),
            "3.1": slice_between(["3.1  Reflections"], ["3.2  Other Suggestions"], body_at),
            "3.2": slice_between(["3.2  Other Suggestions"], ["References"], body_at),
        }
        fig_plan = {
            "1.3": [
                (FIG_EN / "fig3_annotation_flow.png", "Fig. 1-1 Main annotation workflow"),
                (FIG_EN / "fig8_review_flow.png", "Fig. 1-2 Review activity diagram"),
            ],
            "2.1": [
                (FIG_EN / "fig1_architecture.png", "Fig. 2-1 Overall architecture"),
                (FIG_EN / "fig7_modules.png", "Fig. 2-2 My module breakdown"),
            ],
            "2.3": [(FIG_EN / "fig2_usecase.png", "Fig. 2-5 Core use cases")],
            "2.4": [
                (FIG_EN / "fig11_class.png", "Fig. 2-6a Core classes / objects"),
                (FIG_EN / "fig12_state.png", "Fig. 2-6b State machines"),
                (FIG_EN / "fig6_sequence.png", "Fig. 2-8 Save-surgery sequence"),
                (FIG_EN / "fig14_collab.png", "Fig. 2-6c AI predict collaboration"),
                (FIG_EN / "fig5_er.png", "Fig. 2-10 Entity relationships"),
            ],
            "2.5": [(FIG_EN / "fig13_program_flow.png", "Fig. 2-11 Program flow: save surgery ROI")],
            "2.6": [(FIG_EN / "fig10_test_summary.png", "Fig. 2-16 System test summary")],
        }
    return sections, fig_plan


HEADING_MAP_ZH = {
    "1.1": lambda t: t.replace(" ", "").startswith("1.1") and "实训背景" in t and len(t) < 30,
    "1.2": lambda t: t.replace(" ", "").startswith("1.2") and "实训环境" in t and len(t) < 30,
    "1.3": lambda t: t.replace(" ", "").startswith("1.3") and "实训过程" in t and len(t) < 30,
    "2.1": lambda t: t.replace(" ", "").startswith("2.1") and "概述" in t and len(t) < 30,
    "2.2": lambda t: t.replace(" ", "").startswith("2.2") and "相关技术" in t and len(t) < 30,
    "2.3": lambda t: t.replace(" ", "").startswith("2.3") and "系统分析" in t and len(t) < 30,
    "2.4": lambda t: t.replace(" ", "").startswith("2.4") and "系统设计" in t and len(t) < 30,
    "2.5": lambda t: t.replace(" ", "").startswith("2.5") and "系统实现" in t and len(t) < 30,
    "2.6": lambda t: t.replace(" ", "").startswith("2.6") and "系统测试" in t and len(t) < 30,
    "3.1": lambda t: t.replace(" ", "").startswith("3.1") and "实训体会" in t and len(t) < 30,
    "3.2": lambda t: t.replace(" ", "").startswith("3.2") and "其它意见" in t and len(t) < 30,
}


def body_fmt_src(doc):
    # use 说明 first content paragraph as body format reference
    for p in doc.paragraphs:
        if p.text.startswith("实训结束之前") or p.text.startswith("Before the internship"):
            return p
    # fallback: any long justified para
    for p in doc.paragraphs:
        if len(p.text) > 40 and p.runs:
            return p
    return doc.paragraphs[0]


def inject_section_content(doc, key: str, match_fn, items: list, figs: list, fmt_src):
    # Find the LAST matching heading (skip TOC entries that appear earlier)
    heading = None
    for p in doc.paragraphs:
        if match_fn(p.text):
            heading = p
    if heading is None:
        print("WARN: heading not found for", key)
        return

    # Remove empty paragraphs immediately after heading until next non-empty
    nxt = heading._p.getnext()
    while nxt is not None and nxt.tag == qn("w:p"):
        texts = [n.text or "" for n in nxt.iter(qn("w:t"))]
        joined = "".join(texts).strip()
        if joined and re.match(r"^\d+(\.\d+)?", joined.replace("\t", " ").strip()):
            break
        if joined and ("前言" in joined or "实训内容" in joined or "总结" in joined) and len(joined) < 20:
            break
        if not joined:
            nxt2 = nxt.getnext()
            parent = nxt.getparent()
            if parent is not None:
                parent.remove(nxt)
            nxt = nxt2
            continue
        break

    payload = []
    for kind, *rest in items:
        if kind == "text":
            text = rest[0]
            if re.match(r"^\d+(\.\d+)?\s+\S+$", text) and len(text) < 30:
                # keep subsection titles like 2.3.1 as content markers
                pass
            if text in ("说    明", "目  录", "Notes", "Contents"):
                continue
            payload.append(("text", text))

    for path, cap in figs:
        if path.exists():
            payload.append(("fig", path, cap))

    # reverse insert after heading
    for item in reversed(payload):
        if item[0] == "text":
            clone_body_paragraph_after(heading, item[1], fmt_src)
        elif item[0] == "fig":
            _, path, cap = item
            insert_caption_after(heading._p, cap, fmt_src)
            insert_picture_after(heading, path)


def fill_document(lang: str = "zh"):
    base = convert_official_doc()
    work = ROOT / ("_fill_zh.docx" if lang == "zh" else "_fill_en.docx")
    shutil.copy2(base, work)
    doc = Document(str(work))

    if lang == "zh":
        fill_cover_zh(doc)
        setup_header_from_template(doc, "东北大学软件学院", "学生实训总结报告")
        src = ROOT / "_content_src_zh.docx"
        heading_map = HEADING_MAP_ZH
        out = ROOT / "王瑞琦-20236742-软英2301-学生实训总结报告-中文版.docx"
        alias = ROOT / "王瑞琦-20236742-学生实训总结报告-中文翻译版.docx"
    else:
        fill_cover_en(doc)
        setup_header_from_template(
            doc,
            "School of Software, Northeastern University",
            "Student Internship Summary Report",
        )
        src = ROOT / "_content_src_en.docx"
        # For EN, keep Chinese template headings (official form), inject English body under them
        heading_map = HEADING_MAP_ZH
        out = ROOT / "WangRuiqi-20236742-Internship-Summary-Report-EN.docx"
        alias = ROOT / "王瑞琦-20236742-学生实训总结报告-英文版.docx"

    sections, fig_plan = extract_sections_from_generated(src, "zh" if lang == "zh" else "en")
    fmt_src = body_fmt_src(doc)

    # Inject in reverse section order so earlier insertions don't shift search oddly
    # (we re-find headings each time, so order OK either way)
    for key in ["3.2", "3.1", "2.6", "2.5", "2.4", "2.3", "2.2", "2.1", "1.3", "1.2", "1.1"]:
        inject_section_content(
            doc,
            key,
            heading_map[key],
            sections.get(key, []),
            fig_plan.get(key, []),
            fmt_src,
        )

    doc.save(out)
    shutil.copy2(out, alias)
    # also save a clearly named official-filled copy
    if lang == "zh":
        shutil.copy2(out, ROOT / "王瑞琦-20236742-附件05学生实训总结报告-已填写.docx")
    else:
        shutil.copy2(out, ROOT / "WangRuiqi-20236742-Attachment05-Internship-Report-EN.docx")

    # char count
    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    if lang == "zh":
        n = len(re.findall(r"[\u4e00-\u9fff]", text))
        print("ZH filled:", out, "chinese_chars:", n)
    else:
        print("EN filled:", out, "chars:", len(text.replace(" ", "").replace("\n", "")))
    return out


def main():
    fill_document("zh")
    fill_document("en")
    print("DONE")


if __name__ == "__main__":
    main()
